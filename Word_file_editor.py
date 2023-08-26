import docx
import datetime
import tkinter as tk
import os

def main():
    def copy_run_format(source_run, target_run):
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline
        target_run.font.size = source_run.font.size
        target_run.font.name = source_run.font.name

    def search_replace_save():
        Name = Current_Name.get()
        New_name = New_Name.get()
        Address = address.get()
        Relation = relation.get()
        Gaurdian_Name = gaurdian_Name.get()
        wt1_Name = Witness_1.get()
        wt1_Address = w1_Address.get()
        wt1_contact = w1_Contact.get()
        wt2_Name = Witness_2.get()
        wt2_Address = w2_Address.get()
        wt2_contact = w2_Contact.get()
        
        replacements = {
        '@Address': Address,
        'r_elation':Relation,
        'wt1_name': wt1_Name,
        'wt1_address': wt1_Address,
        'wt1_contact': wt1_contact,
        'wt2_name': wt2_Name,
        'wt2_address': wt2_Address,
        'wt2_contact': wt2_contact,
        'ddmmyy': formatted_date
        }

        b_replacements = {
            'current_name': Name,
            'new_name': New_name,
            'gaurdian_name': Gaurdian_Name,
        }
        #declaring input and output file paths
        desktop = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        input_path = 'input.docx'
        output_path = f'output.docx'
        
        doc = docx.Document(input_path)
        
        for para in doc.paragraphs:
            updated_runs = []
            for run in para.runs:
                updated_text = run.text
                for key, value in b_replacements.items():
                    updated_text = updated_text.replace(key, f'<b>{value}</b>')
                for key, value in replacements.items():
                    updated_text = updated_text.replace(key, value)
                
                if '<b>' in updated_text:
                    updated_run = para.add_run()
                    copy_run_format(run, updated_run)
                    updated_run.add_text(updated_text.replace('<b>', '').replace('</b>', ''))
                else:
                    updated_run = para.add_run(updated_text)
                    copy_run_format(run, updated_run)
                
                updated_runs.append(updated_run)
            
            para.clear()
            for run in updated_runs:
                para._element.append(run._element)
        
        doc.save(output_path)
        label_output2.config(text=str(output_path))


    def format_date(date):
        day = date.day
        suffix = "th" if 11 <= day <= 13 else {1: "st", 2: "nd", 3: "rd"}.get(day % 10, "th")
        formatted_date = date.strftime(f"%d{suffix} day of %B %Y")
        return formatted_date


    current_date = datetime.datetime.today()
    formatted_date = str(format_date(current_date))

    # Create the main application window
    global root
    root = tk.Tk()
    root.title("Application Name")
    root.iconbitmap("logo.ico")

    # Change background color to black
    root.configure(bg="black")

    # Create input fields and labels
    label_fg_color = "white"  # Label text color
    input_bg_color = "black"  # Input field background color

    # Increase size by 10%
    font_size = 14  # Base font size
    font_size_increase = int(font_size * 0.1)  # 10% increase

    label_font = ("Helvetica", font_size + font_size_increase, "bold")
    input_font = ("Helvetica", font_size + font_size_increase)

    label1 = tk.Label(root, text="Current Name:", fg=label_fg_color, bg="black", font=label_font)
    label1.grid(row=0, column=0, padx=15, pady=8, sticky="e")
    Current_Name = tk.Entry(root, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    Current_Name.grid(row=0, column=1, padx=15, pady=8)

    label2 = tk.Label(root, text="New Name:", fg=label_fg_color, bg="black", font=label_font)
    label2.grid(row=1, column=0, padx=15, pady=8, sticky="e")
    New_Name = tk.Entry(root, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    New_Name.grid(row=1, column=1, padx=15, pady=8)

    label3 = tk.Label(root, text="Address:", fg=label_fg_color, bg="black", font=label_font)
    label3.grid(row=2, column=0, padx=15, pady=8, sticky="e")
    address = tk.Entry(root,width=20, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    address.grid(row=2, column=1, padx=15, pady=8)
    
    label4 = tk.Label(root, text="Relation (S/O):", fg=label_fg_color, bg="black", font=label_font)
    label4.grid(row=3, column=0, padx=15, pady=8, sticky="e")
    relation = tk.Entry(root, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    relation.grid(row=3, column=1, padx=15, pady=8)

    label5 = tk.Label(root, text="Guardian Name:", fg=label_fg_color, bg="black", font=label_font)
    label5.grid(row=4, column=0, padx=15, pady=8, sticky="e")
    gaurdian_Name = tk.Entry(root, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    gaurdian_Name.grid(row=4, column=1, padx=15, pady=8)

    label6 = tk.Label(root, text="Witness 1:", fg=label_fg_color, bg="black", font=label_font)
    label6.grid(row=5, column=0, padx=15, pady=8, sticky="e")
    Witness_1 = tk.Entry(root, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    Witness_1.grid(row=5, column=1, padx=15, pady=8)

    label7 = tk.Label(root, text="Witness 1 Address:", fg=label_fg_color, bg="black", font=label_font)
    label7.grid(row=6, column=0, padx=15, pady=8, sticky="e")
    w1_Address = tk.Entry(root, width=20, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    w1_Address.grid(row=6, column=1, padx=15, pady=8)

    label8 = tk.Label(root, text="Witness 1 Contact:", fg=label_fg_color, bg="black", font=label_font)
    label8.grid(row=7, column=0, padx=15, pady=8, sticky="e")
    w1_Contact = tk.Entry(root, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    w1_Contact.grid(row=7, column=1, padx=15, pady=8)

    label9 = tk.Label(root, text="Witness 2:", fg=label_fg_color, bg="black", font=label_font)
    label9.grid(row=8, column=0, padx=15, pady=8, sticky="e")
    Witness_2 = tk.Entry(root, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    Witness_2.grid(row=8, column=1, padx=15, pady=8)

    label10 = tk.Label(root, text="Witness 2 Address:", fg=label_fg_color, bg="black", font=label_font)
    label10.grid(row=9, column=0, padx=15, pady=8, sticky="e")
    w2_Address = tk.Entry(root, width=20, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    w2_Address.grid(row=9, column=1, padx=15, pady=8)

    label11 = tk.Label(root, text="Witness 2 Contact:", fg=label_fg_color, bg="black", font=label_font)
    label11.grid(row=10, column=0, padx=15, pady=8, sticky="e")
    w2_Contact = tk.Entry(root, bg=input_bg_color, fg=label_fg_color, insertbackground="white", font=input_font)
    w2_Contact.grid(row=10, column=1, padx=15, pady=8)

    # Continue adding more labels and input fields...
    
    label_output2 = tk.Label(root, text="File Status",font=("Arial", 17))
    
    label_output2.grid(row=11,columnspan=2, padx=5, pady=5)

    # Create a calculate button
    calculate_button = tk.Button(root, text="Generate", command=search_replace_save, fg="green", bg="black", font=label_font)
    calculate_button.grid(row=12, column=1, padx=15, pady=15)

    # Create a restart button
    restart_button = tk.Button(root, text="Restart", command=refresh, fg="red", bg="black", font=label_font)
    restart_button.grid(row=12, column=0, padx=15, pady=15)

    # Start the Tkinter main loop
    root.mainloop()

if __name__ == '__main__':
    def refresh():
        root.destroy()
        main()

main()
